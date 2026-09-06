"""Convierte todos los .docx de docs/ a .md en docs/sii/"""
from pathlib import Path
from docx import Document

DOCS_DIR = Path(__file__).resolve().parent / "docs"
OUT_DIR = DOCS_DIR / "sii"
OUT_DIR.mkdir(exist_ok=True)

def docx_to_md(path: Path) -> str:
    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        elif "list" in style or text.startswith(("-", "•", "*")):
            lines.append(f"- {text.lstrip('-•* ')}")
        else:
            lines.append(text)
    # Tables
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table.rows[1:]:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n".join(lines)

docx_files = list(DOCS_DIR.glob("*.docx"))
print(f"Convirtiendo {len(docx_files)} archivos...")
for f in docx_files:
    md_content = docx_to_md(f)
    out_path = OUT_DIR / (f.stem + ".md")
    out_path.write_text(md_content, encoding="utf-8")
    print(f"  OK: {out_path.name} ({len(md_content)} chars)")

print("Conversión completada.")
